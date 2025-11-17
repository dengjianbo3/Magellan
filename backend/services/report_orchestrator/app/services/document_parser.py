"""
Document Parser Service
文档解析服务

Extracts text content from various document formats (PDF, DOCX, TXT).
从各种文档格式（PDF、DOCX、TXT）中提取文本内容。
"""

from typing import List, Dict, Any
import os
from pypdf import PdfReader
from docx import Document
import re


class DocumentParser:
    """Service for parsing various document formats"""

    @staticmethod
    def parse_pdf(file_path: str) -> Dict[str, Any]:
        """
        Parse PDF file and extract text

        Args:
            file_path: Path to PDF file

        Returns:
            Dictionary with text content and metadata
        """
        try:
            reader = PdfReader(file_path)

            # Extract metadata
            metadata = {
                "num_pages": len(reader.pages),
                "file_type": "pdf",
                "file_name": os.path.basename(file_path)
            }

            # Try to extract PDF metadata
            if reader.metadata:
                if reader.metadata.title:
                    metadata["title"] = reader.metadata.title
                if reader.metadata.author:
                    metadata["author"] = reader.metadata.author
                if reader.metadata.creation_date:
                    metadata["creation_date"] = str(reader.metadata.creation_date)

            # Extract text from all pages
            full_text = ""
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        full_text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                except Exception as e:
                    print(f"[DocumentParser] Error extracting page {page_num}: {e}")
                    continue

            return {
                "text": full_text.strip(),
                "metadata": metadata,
                "success": True
            }

        except Exception as e:
            print(f"[DocumentParser] Error parsing PDF: {e}")
            return {
                "text": "",
                "metadata": {"error": str(e), "file_type": "pdf"},
                "success": False
            }

    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        """
        Parse Word document and extract text

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with text content and metadata
        """
        try:
            doc = Document(file_path)

            # Extract metadata
            metadata = {
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
                "file_type": "docx",
                "file_name": os.path.basename(file_path)
            }

            # Try to extract document properties
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.created:
                    metadata["creation_date"] = str(core_props.created)
            except Exception:
                pass

            # Extract text from paragraphs
            paragraphs_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs_text.append(para.text)

            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        tables_text.append(row_text)

            # Combine all text
            full_text = "\n\n".join(paragraphs_text)
            if tables_text:
                full_text += "\n\n--- Tables ---\n" + "\n".join(tables_text)

            return {
                "text": full_text.strip(),
                "metadata": metadata,
                "success": True
            }

        except Exception as e:
            print(f"[DocumentParser] Error parsing DOCX: {e}")
            return {
                "text": "",
                "metadata": {"error": str(e), "file_type": "docx"},
                "success": False
            }

    @staticmethod
    def parse_txt(file_path: str) -> Dict[str, Any]:
        """
        Parse plain text file

        Args:
            file_path: Path to TXT file

        Returns:
            Dictionary with text content and metadata
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            text = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break  # Success
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise ValueError("Could not decode file with any supported encoding")

            metadata = {
                "file_type": "txt",
                "file_name": os.path.basename(file_path),
                "num_lines": len(text.split('\n')),
                "num_chars": len(text)
            }

            return {
                "text": text.strip(),
                "metadata": metadata,
                "success": True
            }

        except Exception as e:
            print(f"[DocumentParser] Error parsing TXT: {e}")
            return {
                "text": "",
                "metadata": {"error": str(e), "file_type": "txt"},
                "success": False
            }

    @staticmethod
    def parse_document(file_path: str) -> Dict[str, Any]:
        """
        Auto-detect file type and parse document

        Args:
            file_path: Path to document file

        Returns:
            Dictionary with text content and metadata
        """
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if ext == '.pdf':
            return DocumentParser.parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return DocumentParser.parse_docx(file_path)
        elif ext == '.txt':
            return DocumentParser.parse_txt(file_path)
        else:
            return {
                "text": "",
                "metadata": {"error": f"Unsupported file type: {ext}"},
                "success": False
            }

    @staticmethod
    def chunk_text(
        text: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ) -> List[str]:
        """
        Split text into overlapping chunks for better retrieval

        Args:
            text: Input text
            chunk_size: Target size of each chunk (in words)
            chunk_overlap: Number of overlapping words between chunks

        Returns:
            List of text chunks
        """
        # Split by sentences/paragraphs first
        paragraphs = re.split(r'\n\n+', text)

        chunks = []
        current_chunk = []
        current_size = 0

        for para in paragraphs:
            para_words = para.split()
            para_size = len(para_words)

            if current_size + para_size <= chunk_size:
                # Add to current chunk
                current_chunk.append(para)
                current_size += para_size
            else:
                # Save current chunk and start new one
                if current_chunk:
                    chunks.append(" ".join(current_chunk))

                # Start new chunk with overlap
                if current_chunk and chunk_overlap > 0:
                    # Take last few words from previous chunk for overlap
                    overlap_text = " ".join(current_chunk[-1].split()[-chunk_overlap:])
                    current_chunk = [overlap_text, para]
                    current_size = chunk_overlap + para_size
                else:
                    current_chunk = [para]
                    current_size = para_size

        # Add final chunk
        if current_chunk:
            chunks.append(" ".join(current_chunk))

        return [c for c in chunks if c.strip()]  # Filter empty chunks
